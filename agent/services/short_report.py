"""
Short SEO Potential Analysis report — focused, 5-section format.

Designed for sales hand-off: one verdict, the data the analyst would highlight
in a call, and concrete actions. Reuses styling helpers from `word_report.py`.

Sections:
  1. Cover + verdict
  2. Traffic Comparison (client vs competitors, single table)
  3. Top Page-Level Traffic Opportunities (Big Wins)
  4. Keyword-Level Comparison + Rankings (merged)
  5. On-Page Content Check (footer, headings, schema, content depth)
  6. Technical Health (HTTPS, sitemap, robots, etc.)
"""

from io import BytesIO
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor

# Reuse styling helpers from the long report
from agent.services.word_report import (
    _BRAND, _BRAND_DARK, _BRAND_LITE, _WHITE, _GREY, _RED, _AMBER, _GREEN,
    _BRAND_RGB, _BRAND_DARK_RGB, _WHITE_RGB, _DARK_RGB, _GREY_RGB,
    _set_cell_bg, _set_cell_border, _set_row_height, _para_space,
    _add_horizontal_rule, _cover_page, _section_heading, _metric_cards, _add_table,
    _add_footer,
)


def _verdict_color(verdict: str) -> str:
    return {"HIGH": _GREEN, "MEDIUM": _AMBER, "LOW": _RED}.get(verdict.upper(), _GREY)


_OK_RGB    = RGBColor(0x05, 0x96, 0x69)   # green
_WARN_RGB  = RGBColor(0xD9, 0x77, 0x06)   # amber


def _add_finding_line(doc, text: str, *, ok: bool = True):
    """Inline ✓/⚠ list item."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.4)
    icon_run = p.add_run("✓ " if ok else "⚠ ")
    icon_run.bold = True
    icon_run.font.color.rgb = _OK_RGB if ok else _WARN_RGB
    body = p.add_run(text)
    body.font.size = Pt(10)


def generate_short_report(
    *,
    client_url: str,
    client_name: str,
    niche: str,
    location: str,
    ai_result: dict,         # {"potential": ..., "summary": ...}
    gap,                     # GapAnalysis instance — for big_wins, traffic_ratio, notes
    comp_traffic: dict,      # {comp_name: traffic}
    client_total_traffic: int,
    keyword_rank_comparison: pd.DataFrame,
    page_traffic_comparison: pd.DataFrame,
    page_audit: dict | None = None,   # output of page_content_audit.audit()
    site_result=None,                  # SEOAnalyzer().analyze() output
) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.4)
        section.bottom_margin = Cm(1.6)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    date_str  = datetime.now().strftime("%d %B %Y")
    potential = ai_result.get("potential", "MEDIUM")
    summary   = ai_result.get("summary", "")

    _add_footer(doc, date_str)

    # ── 1. Cover ──────────────────────────────────────────────────────────────
    _cover_page(doc, client_url, niche, location, potential, summary, date_str)

    # ── Competitor mis-alignment warning (loud — first thing after cover) ──
    misalign = getattr(gap, "competitor_misalignment", "") or ""
    if misalign:
        wp = doc.add_paragraph()
        wp.paragraph_format.space_before = Pt(8)
        wp.paragraph_format.space_after  = Pt(4)
        wr = wp.add_run("⚠ COMPETITOR MIS-ALIGNMENT DETECTED")
        wr.bold = True
        wr.font.size = Pt(13)
        wr.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)   # red

        wp2 = doc.add_paragraph()
        wp2.paragraph_format.space_after = Pt(6)
        wp2.paragraph_format.left_indent = Cm(0.4)
        wr2 = wp2.add_run(misalign)
        wr2.font.size = Pt(10)
        wr2.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
        _add_horizontal_rule(doc)

    # ── 2. Traffic Comparison ─────────────────────────────────────────────────
    _section_heading(doc, "1.  Traffic Comparison",
                     "Where the client stands vs the competitors selected for analysis")

    rows = []
    rows.append({
        "Website":  f"★ {client_name}",
        "Role":     "CLIENT",
        "Non-Brand Traffic / mo":  f"{client_total_traffic:,}",
        "Gap vs Client":  "—",
    })
    for d, t in sorted(comp_traffic.items(), key=lambda x: -x[1]):
        gap_val = t - client_total_traffic
        ratio = (t / client_total_traffic) if client_total_traffic > 0 else 0
        rows.append({
            "Website":  d,
            "Role":     "Competitor",
            "Non-Brand Traffic / mo":  f"{t:,}",
            "Gap vs Client":  f"{ratio:.1f}× ({'+' if gap_val > 0 else ''}{gap_val:,})",
        })
    _add_table(doc, pd.DataFrame(rows))

    # One-line headline + REALISTIC achievable target (matches Big Wins below)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    avg_comp = sum(comp_traffic.values()) / len(comp_traffic) if comp_traffic else 0
    achievable = getattr(gap, "achievable_top10_traffic", 0)
    if avg_comp > 0 and client_total_traffic > 0:
        ratio = avg_comp / client_total_traffic
        run = p.add_run(
            f"📊 Competitors average {ratio:.1f}× the client's TOTAL traffic. "
            f"But only a fraction of that is reachable from pages relevant to the client's catalog."
        )
        run.font.size = Pt(11)
        run.italic = True
        run.font.color.rgb = _BRAND_DARK_RGB

    # The realistic target — what closing the top-10 relevant gaps gets you
    if achievable > 0:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        run2 = p2.add_run(
            f"🎯 Realistic short-term target: capture ~{achievable:,} clicks/month by ranking "
            f"on the top 10 RELEVANT gap keywords (see page-level opportunities below). "
            f"This is what you can defensibly project in a 6–12 month engagement."
        )
        run2.font.size = Pt(11)
        run2.bold = True
        from docx.shared import RGBColor as _RGB
        run2.font.color.rgb = _RGB(0x05, 0x96, 0x69)   # green

    # ── 3. Top Page-Level Traffic Opportunities (Big Wins) ────────────────────
    if gap.big_wins:
        _section_heading(doc, "2.  Top Page-Level Traffic Opportunities",
                         "Pages to build/optimise first — ranked by traffic competitors actually capture")
        for i, w in enumerate(gap.big_wins, 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(f"#{i}  {w.keyword}")
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = _BRAND_DARK_RGB

            stats = doc.add_paragraph()
            stats.paragraph_format.space_after = Pt(2)
            sr = stats.add_run(
                f"Volume: {w.volume:,}/mo  ·  "
                f"{w.competitor} captures {w.competitor_traffic:,} clicks/mo  ·  "
                f"Client rank: {w.client_rank}"
            )
            sr.font.size = Pt(10)
            sr.font.color.rgb = _GREY_RGB

            pitch = doc.add_paragraph()
            pitch.paragraph_format.space_after  = Pt(8)
            pitch.paragraph_format.left_indent = Cm(0.5)
            pr = pitch.add_run(f"💡  {w.pitch}")
            pr.font.size = Pt(10)
            pr.italic = True

    # ── 4. Keyword-Level Comparison + Rankings ────────────────────────────────
    if not keyword_rank_comparison.empty:
        _section_heading(doc, "3.  Keyword-Level Traffic Opportunities",
                         "Top high-volume keywords — where each site ranks (NR = not ranking)")
        # Trim to the most relevant 12-15 rows
        display = keyword_rank_comparison.head(15)
        _add_table(doc, display, max_rows=15)

    # ── 5. On-Page Content Check ──────────────────────────────────────────────
    if page_audit:
        _section_heading(doc, "4.  On-Page Content Check",
                         f"Audit of the client's top page: {page_audit.get('url', client_url)[:80]}")

        _metric_cards(doc, [
            ("H1 Tags",            f"{page_audit.get('h1_count', 0)}"),
            ("H2 Tags",            f"{page_audit.get('h2_count', 0)}"),
            ("Above-fold Words",   f"{page_audit.get('above_fold_word_count', 0):,}"),
            ("Footer Words",       f"{page_audit.get('footer_word_count', 0):,}"),
            ("Schema Types",       f"{len(page_audit.get('schema_types', []))}"),
        ])

        # Findings
        issues = page_audit.get("issues", [])
        if not issues:
            _add_finding_line(doc, "All on-page content checks passed. Page is well-structured.", ok=True)
        else:
            for issue in issues[:8]:
                _add_finding_line(doc, issue, ok=False)

        # Schema list
        if page_audit.get("schema_types"):
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Cm(0.4)
            sp.paragraph_format.space_before = Pt(6)
            sr = sp.add_run("Schema markup detected: ")
            sr.bold = True
            sr.font.size = Pt(10)
            sr2 = sp.add_run(", ".join(page_audit["schema_types"]))
            sr2.font.size = Pt(10)

    # ── 6. Technical Health ───────────────────────────────────────────────────
    if site_result:
        _section_heading(doc, "5.  Technical Health",
                         "Foundational SEO signals on the client's homepage")
        sig = site_result.key_signals
        load_time = sig.get("load_time_s")
        load_str = f"{load_time}s" if load_time is not None else "—"

        rows = [
            {"Check": "HTTPS",            "Status": "✓ Enabled" if sig.get("uses_https")        else "⚠ Not enabled"},
            {"Check": "Page load time",   "Status": load_str + (" (slow)" if load_time and load_time > 3 else "")},
            {"Check": "Meta description", "Status": "✓ Present" if sig.get("meta_desc_len")     else "⚠ Missing"},
            {"Check": "Canonical tag",    "Status": "✓ Present" if sig.get("canonical")         else "⚠ Missing"},
            {"Check": "Sitemap.xml",      "Status": "✓ Found"   if sig.get("has_sitemap")       else "⚠ Missing"},
            {"Check": "Robots.txt",       "Status": "✓ Found"   if sig.get("robots_txt_exists") else "⚠ Missing"},
            {"Check": "Mobile viewport",  "Status": "✓ Present" if sig.get("has_viewport")      else "⚠ Missing"},
            {"Check": "Schema markup",    "Status": "✓ Present" if sig.get("has_schema")        else "⚠ Missing"},
            {"Check": "OG tags",          "Status": "✓ Present" if sig.get("has_og")            else "⚠ Missing"},
            {"Check": "Images w/o alt",   "Status": f"{sig.get('img_no_alt', 0)} image(s)"},
        ]
        _add_table(doc, pd.DataFrame(rows))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
