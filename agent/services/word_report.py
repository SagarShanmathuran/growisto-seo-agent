"""Generate a polished Word document SEO Potential Analysis report — Growisto theme."""

from io import BytesIO
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agent.services.strategic_advisor import build_strategic_recommendations

# ── Brand colours ─────────────────────────────────────────────────────────────
_BRAND      = "367588"          # Growisto teal
_BRAND_DARK = "1D3F4A"         # darker teal for headings
_BRAND_LITE = "E8F4F7"         # very light teal for alternate rows
_WHITE      = "FFFFFF"
_DARK       = "1A1A2E"         # near-black for body text
_GREY       = "6B7280"         # muted grey for captions
_RED        = "DC2626"         # high priority
_AMBER      = "D97706"         # medium priority
_GREEN      = "059669"         # low priority / positive

_BRAND_RGB      = RGBColor(0x36, 0x75, 0x88)
_BRAND_DARK_RGB = RGBColor(0x1D, 0x3F, 0x4A)
_WHITE_RGB      = RGBColor(0xFF, 0xFF, 0xFF)
_DARK_RGB       = RGBColor(0x1A, 0x1A, 0x2E)
_GREY_RGB       = RGBColor(0x6B, 0x72, 0x80)


# ── XML helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, color in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if color:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:color"), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_row_height(row, height_cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trPr.append(trHeight)


def _para_space(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


def _add_horizontal_rule(doc, color=_BRAND):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    _para_space(p, before=0, after=60)
    return p


# ── Cover page ─────────────────────────────────────────────────────────────────

def _cover_page(doc, client_url, niche, location, potential, summary, date_str):
    # Brand colour block (simulated via a full-width 1-row table)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, _BRAND)
    _set_row_height(tbl.rows[0], 3.5)

    # Title inside the band
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)

    run = p.add_run("SEO POTENTIAL ANALYSIS")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = _WHITE_RGB
    run.font.name = "Calibri"

    p2 = cell.add_paragraph("REPORT")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.runs[0]
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0xB2, 0xD8, 0xE4)
    r2.font.name = "Calibri"

    doc.add_paragraph()

    # Client info block
    info_tbl = doc.add_table(rows=4, cols=2)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    labels = ["Client Website", "Industry / Niche", "Target Market", "Report Date"]
    values = [client_url, niche, location, date_str]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        lc = info_tbl.rows[i].cells[0]
        vc = info_tbl.rows[i].cells[1]
        _set_cell_bg(lc, _BRAND_DARK)
        _set_cell_bg(vc, _BRAND_LITE if i % 2 == 0 else _WHITE)
        lp = lc.paragraphs[0]
        lp.add_run(lbl).font.color.rgb = _WHITE_RGB
        lp.runs[0].bold = True
        lp.runs[0].font.size = Pt(9)
        vp = vc.paragraphs[0]
        vp.add_run(val).font.size = Pt(9)
        vp.runs[0].font.color.rgb = _DARK_RGB

    doc.add_paragraph()

    # Verdict badge
    verdict_bg   = {"HIGH": _RED, "MEDIUM": _AMBER, "LOW": "2563EB"}.get(potential, _BRAND)
    verdict_text = {"HIGH": "HIGH POTENTIAL", "MEDIUM": "MEDIUM POTENTIAL", "LOW": "LOW POTENTIAL"}.get(potential, potential)
    v_tbl = doc.add_table(rows=1, cols=1)
    v_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    vc = v_tbl.rows[0].cells[0]
    _set_cell_bg(vc, verdict_bg)
    _set_row_height(v_tbl.rows[0], 1.0)
    vp = vc.paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vp.paragraph_format.space_before = Pt(6)
    vr = vp.add_run(f"SEO POTENTIAL:  {verdict_text}")
    vr.bold = True
    vr.font.size = Pt(16)
    vr.font.color.rgb = _WHITE_RGB
    vr.font.name = "Calibri"

    doc.add_paragraph()

    # Summary box
    if summary:
        s_tbl = doc.add_table(rows=1, cols=1)
        s_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        sc = s_tbl.rows[0].cells[0]
        _set_cell_bg(sc, _BRAND_LITE)
        _set_cell_border(sc, top=_BRAND, bottom=_BRAND, left=_BRAND, right=_BRAND)
        sp = sc.paragraphs[0]
        sp.add_run("Executive Summary  ").bold = True
        sp.runs[0].font.color.rgb = _BRAND_DARK_RGB
        sp.runs[0].font.size = Pt(10)
        sp2 = sc.add_paragraph(summary)
        sp2.runs[0].font.size = Pt(9)
        sp2.runs[0].font.color.rgb = _DARK_RGB

    doc.add_page_break()


# ── Section heading ────────────────────────────────────────────────────────────

def _section_heading(doc, text: str, subtitle: str = ""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, _BRAND)
    _set_row_height(tbl.rows[0], 0.75)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _WHITE_RGB
    run.font.name = "Calibri"
    if subtitle:
        p2 = cell.add_paragraph(subtitle)
        p2.runs[0].font.size = Pt(8)
        p2.runs[0].font.color.rgb = RGBColor(0xB2, 0xD8, 0xE4)
    doc.add_paragraph()


# ── Metric cards row ───────────────────────────────────────────────────────────

def _metric_cards(doc, cards: list[tuple[str, str]]):
    """cards = [(label, value), ...]"""
    tbl = doc.add_table(rows=2, cols=len(cards))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (label, value) in enumerate(cards):
        top = tbl.rows[0].cells[i]
        bot = tbl.rows[1].cells[i]
        _set_cell_bg(top, _BRAND)
        _set_cell_bg(bot, _BRAND_LITE)
        tp = top.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run(label.upper())
        tr.bold = True
        tr.font.size = Pt(7)
        tr.font.color.rgb = RGBColor(0xB2, 0xD8, 0xE4)
        bp = bot.paragraphs[0]
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        br = bp.add_run(value)
        br.bold = True
        br.font.size = Pt(13)
        br.font.color.rgb = _BRAND_DARK_RGB
    doc.add_paragraph()


# ── Styled data table ──────────────────────────────────────────────────────────

def _add_table(doc, df: pd.DataFrame, max_rows=30):
    if df.empty:
        p = doc.add_paragraph("No data available.")
        p.runs[0].font.color.rgb = _GREY_RGB
        p.runs[0].font.size = Pt(9)
        return

    df = df.head(max_rows)
    cols = list(df.columns)
    table = doc.add_table(rows=1 + len(df), cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(cols):
        cell = hrow.cells[i]
        _set_cell_bg(cell, _BRAND)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h))
        run.bold = True
        run.font.color.rgb = _WHITE_RGB
        run.font.size = Pt(8)
        run.font.name = "Calibri"

    # Data rows
    for r_idx, (_, row) in enumerate(df.iterrows(), start=1):
        bg = _BRAND_LITE if r_idx % 2 == 0 else _WHITE
        for c_idx, col in enumerate(cols):
            cell = table.rows[r_idx].cells[c_idx]
            _set_cell_bg(cell, bg)
            val = row[col]
            cell.text = str(val) if pd.notna(val) else ""
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].font.size = Pt(8)
                p.runs[0].font.color.rgb = _DARK_RGB

    doc.add_paragraph()


# ── Priority badge paragraph ───────────────────────────────────────────────────

def _priority_badge(doc, title: str, priority: str, rationale: str, actions: list[str]):
    badge_bg   = {"High": _RED, "Medium": _AMBER, "Low": _GREEN}.get(priority, _BRAND)
    badge_text = {"High": "HIGH PRIORITY", "Medium": "MEDIUM PRIORITY", "Low": "LOW PRIORITY"}.get(priority, priority)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # narrow badge column
    tbl.columns[0].width = Cm(3)
    badge_cell = tbl.rows[0].cells[0]
    title_cell = tbl.rows[0].cells[1]
    _set_cell_bg(badge_cell, badge_bg)
    _set_cell_bg(title_cell, _BRAND_DARK)
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bp.paragraph_format.space_before = Pt(6)
    br = bp.add_run(badge_text)
    br.bold = True
    br.font.size = Pt(7)
    br.font.color.rgb = _WHITE_RGB

    tp = title_cell.paragraphs[0]
    tp.paragraph_format.space_before = Pt(4)
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(11)
    tr.font.color.rgb = _WHITE_RGB
    tr.font.name = "Calibri"

    # rationale
    rp = doc.add_paragraph(rationale)
    rp.runs[0].font.size = Pt(9)
    rp.runs[0].font.color.rgb = _GREY_RGB
    rp.runs[0].italic = True
    _para_space(rp, before=40, after=40)

    # actions
    for action in actions:
        ap = doc.add_paragraph(style="List Bullet")
        # strip markdown bold markers for Word
        clean = action.replace("**", "")
        ar = ap.add_run(clean)
        ar.font.size = Pt(9)
        ar.font.color.rgb = _DARK_RGB

    doc.add_paragraph()


# ── Footer ─────────────────────────────────────────────────────────────────────

def _add_footer(doc, date_str):
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(f"Growisto SEO Potential Analyzer  |  Confidential  |  {date_str}")
        fr.font.size = Pt(7)
        fr.font.color.rgb = _GREY_RGB


# ── Main generator ─────────────────────────────────────────────────────────────

def generate_word_report(
    client_url: str,
    niche: str,
    location: str,
    currency_symbol: str,
    ai_result: dict,
    roi: dict,
    comp_traffic: dict,
    ahrefs: dict,
    site_result=None,
    roi_scenarios: dict | None = None,
    include_roi: bool = False,    # Hide revenue/ROI from client report by default (internal-only)
    max_keyword_gaps: int = 15,
    max_low_hanging: int = 10,
    max_top_pages: int = 10,
) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    date_str    = datetime.now().strftime("%d %B %Y")
    potential   = ai_result.get("potential", "MEDIUM")
    summary     = ai_result.get("summary", "")
    all_gaps    = ahrefs.get("keyword_gaps",     pd.DataFrame())
    lh_df       = ahrefs.get("low_hanging_fruit", pd.DataFrame())
    tp_df       = ahrefs.get("top_comp_pages",    pd.DataFrame())
    kw_rank_cmp = ahrefs.get("keyword_rank_comparison", pd.DataFrame())
    pg_traf_cmp = ahrefs.get("page_traffic_comparison", pd.DataFrame())
    client_trf  = ahrefs.get("client_total_traffic", 0)
    total_comp  = sum(comp_traffic.values()) if comp_traffic else 0

    _add_footer(doc, date_str)

    # ── Cover page ────────────────────────────────────────────────────────────
    _cover_page(doc, client_url, niche, location, potential, summary, date_str)

    # ── ROI Summary (internal-only by default; include_roi=True to show) ────
    if include_roi:
        cur_trf  = roi.get("client_current_traffic", 0)
        tgt_trf  = roi.get("target_traffic", 0)
        aov_val  = roi.get("aov", 0)
        cvr_val  = roi.get("conversion_rate", 0)
        cost_val = roi.get("monthly_seo_cost", 0)

        _section_heading(
            doc,
            "SEO ROI Estimate (Internal)",
            f"Incremental-traffic model · {cvr_val:.1%} CVR · AOV {currency_symbol}{aov_val:,.0f} · Retainer {currency_symbol}{cost_val:,.0f}/mo",
        )
        viability = roi.get("viability", "—").upper()
        _metric_cards(doc, [
            ("Current Traffic",        f"{cur_trf:,} / mo"),
            ("Target Traffic",         f"{tgt_trf:,} / mo"),
            ("Incremental",            f"+{roi.get('incremental_traffic', 0):,} / mo"),
            ("Est. Monthly Revenue",   f"{currency_symbol}{roi.get('monthly_revenue', 0):,.0f}"),
            ("ROI",                    f"{roi.get('roi_multiple', 0)}x  ·  {viability}"),
        ])
        if roi_scenarios and all(k in roi_scenarios for k in ("conservative", "realistic", "aggressive")):
            scenario_rows = []
            for label, key in [("Conservative (lowest comp)", "conservative"),
                               ("Realistic (avg comp)",       "realistic"),
                               ("Aggressive (top comp)",      "aggressive")]:
                sc = roi_scenarios[key]
                scenario_rows.append({
                    "Scenario":      label,
                    "Target / mo":   f"{sc['target_traffic']:,}",
                    "Incremental":   f"+{sc['incremental_traffic']:,}",
                    "Revenue / mo":  f"{currency_symbol}{sc['monthly_revenue']:,.0f}",
                    "ROI":           f"{sc['roi_multiple']}x",
                    "Verdict":       sc["viability"].upper(),
                })
            _add_table(doc, pd.DataFrame(scenario_rows))

    # ── Competitor Traffic ─────────────────────────────────────────────────────
    if comp_traffic:
        _section_heading(doc, "Competitor Traffic Benchmark")
        rows = []
        if client_trf:
            rows.append({"": "★ CLIENT", "Website": client_url, "Est. Organic Traffic / Month": f"{client_trf:,}"})
        for d, t in sorted(comp_traffic.items(), key=lambda x: -x[1]):
            gap = t - client_trf
            rows.append({"": "", "Website": d,
                         "Est. Organic Traffic / Month": f"{t:,}",
                         "Traffic Gap": f"+{gap:,}" if gap > 0 else f"{gap:,}"})
        _add_table(doc, pd.DataFrame(rows))

    # ── Page-Level Traffic Comparison ─────────────────────────────────────────
    if not pg_traf_cmp.empty:
        _section_heading(
            doc,
            "Page-Level Traffic Comparison",
            "Same category, different sites — how much traffic each page wins",
        )
        _add_table(doc, pg_traf_cmp, max_rows=12)

    # ── Keyword-Level Ranking Comparison ──────────────────────────────────────
    if not kw_rank_cmp.empty:
        _section_heading(
            doc,
            "Keyword Ranking Comparison",
            "Where each site ranks for top high-volume keywords (NR = not ranking)",
        )
        _add_table(doc, kw_rank_cmp, max_rows=15)

    # ── Keyword Gaps ──────────────────────────────────────────────────────────
    if not all_gaps.empty:
        _section_heading(
            doc,
            f"Top {min(max_keyword_gaps, len(all_gaps))} Page-Level Traffic Opportunities",
            "Pages where competitors rank top 20 and the client doesn't — sorted by traffic the competitor actually wins",
        )
        display = all_gaps[[c for c in [
            "keyword", "search_volume", "competitor", "competitor_position",
            "competitor_traffic", "client_position",
        ] if c in all_gaps.columns]].copy()
        # Sort by competitor_traffic if available (real traffic > theoretical volume)
        if "competitor_traffic" in display.columns:
            display = display.sort_values("competitor_traffic", ascending=False)
        display = display.head(max_keyword_gaps)
        display.columns = [c.replace("_", " ").title() for c in display.columns]
        _add_table(doc, display, max_rows=max_keyword_gaps)

    # ── Low-Hanging Fruit ─────────────────────────────────────────────────────
    if not lh_df.empty:
        _section_heading(doc, f"Low-Hanging Fruit  —  Top {min(max_low_hanging, len(lh_df))} Quick Wins",
                         "Client ranks positions 11–30 on these — small push to page 1 captures meaningful traffic")
        display = lh_df[[c for c in [
            "keyword", "current_position", "search_volume", "traffic_if_top5",
        ] if c in lh_df.columns]].head(max_low_hanging).copy()
        display.columns = [c.replace("_", " ").title() for c in display.columns]
        _add_table(doc, display, max_rows=max_low_hanging)

    # ── Top Competitor Pages ───────────────────────────────────────────────────
    if not tp_df.empty:
        _section_heading(doc, f"Top {min(max_top_pages, len(tp_df))} Competitor Pages by Traffic",
                         "High-traffic pages to replicate or out-rank")
        display = tp_df[[c for c in [
            "competitor", "url", "traffic", "top_keyword"
        ] if c in tp_df.columns]].head(max_top_pages).copy()
        display.columns = [c.replace("_", " ").title() for c in display.columns]
        _add_table(doc, display, max_rows=max_top_pages)

    # ── Strategic Recommendations (top 3 priorities only for compactness) ──
    doc.add_page_break()
    _section_heading(doc, "Strategic Recommendations",
                     "Top 3 priorities — full action plan available on request")
    strat_recs = build_strategic_recommendations(ahrefs, site_result)[:3]
    for rec in strat_recs:
        _priority_badge(doc, rec["title"], rec["priority"], rec["rationale"], rec["actions"])

    # ── On-Page Health ─────────────────────────────────────────────────────────
    if site_result:
        _section_heading(doc, "On-Page Health Check")
        sig = site_result.key_signals
        _metric_cards(doc, [
            ("Health Score",    f"{site_result.health_score:.0f} / 100"),
            ("Potential Score", f"{site_result.potential_score:.0f} / 100"),
            ("Checks Passed",   str(sig.get("checks_passed", 0))),
            ("Checks Failed",   str(sig.get("checks_failed", 0))),
        ])
        audit_rows = [
            {"Check": "HTTPS",            "Status": "✅ Enabled"  if sig.get("uses_https")    else "❌ Not enabled"},
            {"Check": "Page Load Time",   "Status": f"{sig.get('load_time_s', '?')}s"                              },
            {"Check": "Meta Description", "Status": "✅ Present"  if sig.get("meta_desc_len") else "⚠️ Missing"   },
            {"Check": "H1 Tags",          "Status": f"{sig.get('h1_count', 0)} found (1 expected)"                },
            {"Check": "Sitemap.xml",      "Status": "✅ Found"    if sig.get("has_sitemap")   else "⚠️ Missing"   },
            {"Check": "Schema Markup",    "Status": "✅ Found"    if sig.get("has_schema")    else "⚠️ Missing"   },
            {"Check": "Images w/o Alt",   "Status": f"{sig.get('img_no_alt', 0)} image(s) missing alt text"       },
        ]
        _add_table(doc, pd.DataFrame(audit_rows))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
