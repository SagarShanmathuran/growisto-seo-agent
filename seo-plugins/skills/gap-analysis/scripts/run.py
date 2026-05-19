"""
gap-analysis runner. Loads Ahrefs keyword CSVs for client + competitors, applies
intent + URL filters, computes gap keywords, picks Big Wins, writes a Word report.

Claude (in the chat) does the keyword-relevance reasoning. This script provides
two modes:

  - Default (no --in-scope-keywords): writes JSON listing all candidate gap kws
    for Claude to classify in chat.
  - --finalize with --in-scope-keywords PATH: regenerates the report using
    only the kws Claude marked in_scope.

Usage:
  python3 run.py --client-name VINOD --client-url vinodcookware.com \\
      --client-csv /path/client.csv \\
      --competitor "Milton:/path/milton.csv" \\
      --competitor "Borosil:/path/borosil.csv" \\
      [--business-model b2c_ecommerce] \\
      [--niche cookware] \\
      [--output-dir /tmp]
      [--in-scope-keywords /tmp/in-scope.json --finalize]
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd


_NON_SERVICE_URL_RE = re.compile(
    r"/(blog|blogs|glossary|insights|resources|guide|guides|news|articles|"
    r"learn|help|support|careers|press|about|library|webinar|ebook|whitepaper)/",
    re.I,
)


# ── CSV loader (handles UTF-16 tab AND UTF-8 comma) ──────────────────────────

def _read_ahrefs_csv(path: str | Path) -> pd.DataFrame:
    for enc, sep in [("utf-16", "\t"), ("utf-8", ","), ("utf-8", "\t"), ("latin-1", ",")]:
        try:
            df = pd.read_csv(path, encoding=enc, sep=sep)
            if df.shape[1] > 1:
                return df
        except (UnicodeDecodeError, UnicodeError, pd.errors.ParserError):
            continue
    raise ValueError(f"Could not parse {path}")


def _is_top_pages_format(df: pd.DataFrame) -> bool:
    """Detect Ahrefs Top Pages export: has a 'Top keyword' column (or variant)."""
    cols_lower = {c.lower().strip() for c in df.columns}
    return any(c == "top keyword" or c.startswith("top keyword ") or c.startswith("top keyword:")
               for c in cols_lower)


def _normalise_top_pages_format(df: pd.DataFrame) -> pd.DataFrame:
    """Map Ahrefs Top Pages columns to the keyword-schema run.py expects.

    Each row in Top Pages export = one URL + its top-ranking keyword. We treat
    the top keyword as if it were the row's primary keyword for gap analysis.
    """
    rename: dict[str, str] = {}
    intent_combined_col = None

    for col in df.columns:
        cl = col.lower().strip()
        # The keyword itself
        if cl == "top keyword":
            rename[col] = "Keyword"
        # Volume variants
        elif cl in ("top keyword: volume", "top keyword volume", "volume"):
            rename[col] = "Volume"
        # Position variants
        elif cl in ("top keyword: position", "top keyword position", "position"):
            rename[col] = "Current position"
        # Traffic — Top Pages CSV usually has page-level traffic; use that as
        # the keyword's traffic (the top kw drives most of a page's traffic)
        elif cl in ("current traffic", "organic traffic", "traffic"):
            rename[col] = "Current organic traffic"
        # URL
        elif cl in ("current url", "url", "page url"):
            rename[col] = "Current URL"
        # Intent — Ahrefs sometimes ships a single comma-joined column
        elif cl in ("top keyword: intents", "top keyword intents",
                    "intents", "intent"):
            intent_combined_col = col

    out = df.rename(columns=rename) if rename else df.copy()

    # If intent came as a single combined column, expand into boolean flags
    if intent_combined_col and "Commercial" not in out.columns:
        intents = out[intent_combined_col].astype(str).str.lower()
        out["Commercial"]    = intents.str.contains("commercial",    na=False)
        out["Transactional"] = intents.str.contains("transactional", na=False)
        out["Branded"]       = intents.str.contains("branded",       na=False)

    # Some Top Pages exports omit Branded entirely. We need to be conservative:
    # mark all as not-branded if missing, so we don't accidentally filter
    # everything out downstream.
    if "Branded" not in out.columns:
        out["Branded"] = False

    # Same for Commercial/Transactional — if missing, mark all True so the
    # intent filter doesn't drop everything (Top Pages export is intent-light
    # in older formats).
    if "Commercial" not in out.columns:
        out["Commercial"] = True
    if "Transactional" not in out.columns:
        out["Transactional"] = True

    return out


def _normalise_kw_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Map any Ahrefs export (Organic Keywords or Top Pages, snapshot or comparison)
    onto the canonical keyword-schema run.py uses downstream."""
    # If it's a Top Pages export, hand off to the dedicated normaliser first
    if _is_top_pages_format(df):
        df = _normalise_top_pages_format(df)

    # Then apply standard snapshot→comparison aliases
    rename = {}
    for canonical, alts in {
        "Current organic traffic":  ["Organic traffic"],
        "Current position":         ["Position"],
        "Current URL":              ["URL"],
        "Organic traffic change":   ["Traffic change"],
    }.items():
        if canonical not in df.columns:
            for alt in alts:
                if alt in df.columns:
                    rename[alt] = canonical
                    break
    return df.rename(columns=rename) if rename else df


def load_keywords(path: str | Path) -> pd.DataFrame:
    raw = _read_ahrefs_csv(path)
    is_top_pages = _is_top_pages_format(raw)
    df = _normalise_kw_columns(raw)
    required = ["Keyword", "Branded", "Volume", "Current position", "Current organic traffic"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise ValueError(
            f"CSV {path} missing columns after normalising: {missing}. "
            f"Got: {list(df.columns)}. "
            f"Detected format: {'top_pages' if is_top_pages else 'organic_keywords'}"
        )
    df["Current position"]        = pd.to_numeric(df["Current position"], errors="coerce")
    df["Volume"]                  = pd.to_numeric(df["Volume"], errors="coerce").fillna(0)
    df["Current organic traffic"] = pd.to_numeric(df["Current organic traffic"], errors="coerce").fillna(0)
    # Drop rows where Keyword is null/empty (Top Pages export sometimes has
    # blank top-keyword rows for image/non-textual results)
    df = df[df["Keyword"].astype(str).str.strip().str.len() > 0].copy()
    df.attrs["source_format"] = "top_pages" if is_top_pages else "organic_keywords"
    return df


# ── Intent + URL filter (from existing agent/) ───────────────────────────────

def _is_b2b(business_model: str) -> bool:
    bm = (business_model or "").lower()
    return bm.startswith("b2b") or bm in ("financial", "saas")


def _is_ecomm(business_model: str) -> bool:
    bm = (business_model or "").lower()
    return bm.startswith("b2c") or bm in ("ecommerce", "retail")


def _filter_intent_and_urls(df: pd.DataFrame, business_model: str) -> pd.DataFrame:
    """Drop info/blog kws + service-page-only filter for B2B/ecomm."""
    if not (_is_b2b(business_model) or _is_ecomm(business_model)):
        return df
    out = df
    if "Commercial" in out.columns and "Transactional" in out.columns:
        out = out[(out["Commercial"] == True) | (out["Transactional"] == True)]
    if "Current URL" in out.columns:
        mask = out["Current URL"].fillna("").astype(str).apply(
            lambda u: not bool(_NON_SERVICE_URL_RE.search(u)))
        out = out[mask]
    return out.copy()


# ── Gap-keyword computation ──────────────────────────────────────────────────

def compute_gap_keywords(client_df: pd.DataFrame,
                         comps: list[tuple[str, pd.DataFrame]],
                         min_gap_volume: int = 500,
                         max_gap_keywords: int = 30) -> list[dict]:
    # Build keyword -> {position, url, traffic} lookup so we can surface the
    # client's current page (if any) AND its actual traffic on each gap kw
    client_lookup: dict[str, dict] = {}
    has_url_col = "Current URL" in client_df.columns
    for _, r in client_df.iterrows():
        client_lookup[r["Keyword"]] = {
            "position": r["Current position"],
            "url":      str(r.get("Current URL", "") or "") if has_url_col else "",
            "traffic":  int(r.get("Current organic traffic", 0) or 0),
        }
    gap_rows: dict[str, dict] = {}
    for name, cnb in comps:
        top10 = cnb[(cnb["Current position"] <= 10) & (cnb["Volume"] >= min_gap_volume)]
        for _, row in top10.iterrows():
            kw = row["Keyword"]
            client_info = client_lookup.get(kw, {})
            client_pos = client_info.get("position")
            client_pos_n = client_pos if pd.notna(client_pos) else 999
            if client_pos_n <= 20:
                continue
            comp_rank = int(row["Current position"])
            existing = gap_rows.get(kw)
            if existing and existing["competitor_rank"] <= comp_rank:
                continue
            gap_rows[kw] = {
                "keyword":            kw,
                "volume":             int(row["Volume"]),
                "best_competitor":    name,
                "competitor_rank":    comp_rank,
                "competitor_traffic": int(row.get("Current organic traffic", 0) or 0),
                "competitor_url":     str(row.get("Current URL", "") or ""),
                "client_rank":        "NR" if client_pos_n == 999 else str(int(client_pos_n)),
                "client_url":         client_info.get("url", "") if client_pos_n < 999 else "",
                "client_traffic":     client_info.get("traffic", 0) if client_pos_n < 999 else 0,
            }
    return sorted(gap_rows.values(), key=lambda x: -x["volume"])[:max_gap_keywords]


# ── Page-level opportunity aggregation ───────────────────────────────────────

def _page_slug(url: str) -> str:
    """Pull a readable label from a URL — last path segment with hyphens→spaces."""
    if not url:
        return "(unknown page)"
    from urllib.parse import urlparse
    try:
        p = urlparse(url)
        segments = [s for s in p.path.strip("/").split("/") if s]
        last = segments[-1] if segments else (p.netloc or "(homepage)")
        # Strip .html / .php / query bits
        last = re.split(r"\.(?:html?|php|aspx)$", last, flags=re.I)[0]
        return last.replace("-", " ").replace("_", " ").strip() or "(page)"
    except Exception:
        return "(page)"


def page_opportunities(gap_kws: list[dict], top_n: int = 5) -> list[dict]:
    """Group gap keywords by competitor page URL. Each group = one page-level
    opportunity. Ranked by total competitor traffic the page captures from
    in-scope kws. Each entry knows whether the client has an existing page
    to optimise or needs to build a new one."""
    pages: dict[str, dict] = {}
    for g in gap_kws:
        url = g.get("competitor_url") or ""
        comp = g.get("best_competitor", "")
        key = f"{comp}||{url}" if url else f"{comp}||{g['keyword']}"
        if key not in pages:
            pages[key] = {
                "competitor":              comp,
                "competitor_url":          url,
                "page_label":              _page_slug(url),
                "keywords":                [],
                "total_search_volume":     0,
                "competitor_traffic_capture": 0,
                "client_traffic_on_cluster": 0,
                "best_competitor_rank":    99,
                "client_existing_pages":   set(),
                "client_has_weak_ranking": False,
            }
        p = pages[key]
        p["keywords"].append({
            "kw":      g["keyword"],
            "vol":     g["volume"],
            "cr":      g["competitor_rank"],
            "client_rank": g.get("client_rank", "NR"),
            "client_traffic": g.get("client_traffic", 0),
        })
        p["total_search_volume"]        += g["volume"]
        p["competitor_traffic_capture"]  += g["competitor_traffic"]
        p["client_traffic_on_cluster"]   += g.get("client_traffic", 0)
        if g["competitor_rank"] < p["best_competitor_rank"]:
            p["best_competitor_rank"] = g["competitor_rank"]
        if g.get("client_rank") and g["client_rank"] != "NR":
            try:
                if 11 <= int(g["client_rank"]) <= 30:
                    p["client_has_weak_ranking"] = True
            except ValueError:
                pass
            if g.get("client_url"):
                p["client_existing_pages"].add(g["client_url"])

    # Tidy + label
    result = []
    for p in pages.values():
        p["keywords"].sort(key=lambda k: -k["vol"])
        p["top_keyword"]    = p["keywords"][0]["kw"] if p["keywords"] else ""
        p["keyword_count"]  = len(p["keywords"])
        p["action_type"]    = (
            "OPTIMISE EXISTING PAGE" if p["client_has_weak_ranking"] or p["client_existing_pages"]
            else "BUILD NEW PAGE"
        )
        p["client_existing_pages"] = sorted(p["client_existing_pages"])
        result.append(p)

    result.sort(key=lambda p: -p["competitor_traffic_capture"])
    return result[:top_n]


def _score_big_win(g: dict) -> float:
    score = g["competitor_traffic"]
    if g["competitor_rank"] > 1: score *= 1.2
    if g["competitor_rank"] > 3: score *= 1.1
    if g["client_rank"] != "NR":
        try:
            if 11 <= int(g["client_rank"]) <= 30:
                score *= 1.3
        except ValueError:
            pass
    return score


def quick_wins_existing_rankings(gap_kws: list[dict], top_n: int = 10) -> list[dict]:
    """Filter gap_kws to kws where the client already ranks pos 11-30. These are
    the fastest wins — page exists, just needs on-page work. Adds a templated
    `recommended_action` string keyed off the current rank."""
    out = []
    for g in gap_kws:
        rank = g.get("client_rank", "NR")
        if rank == "NR":
            continue
        try:
            r = int(rank)
        except (ValueError, TypeError):
            continue
        if not (11 <= r <= 30):
            continue
        if r <= 15:
            action = f"Already page 2 — title tag + internal links should push toward top 10"
        elif r <= 25:
            action = f"Page 2 — content depth + internal links from authority pages needed"
        else:
            action = f"Page 3 — content overhaul + consider consolidating with stronger page"
        out.append({**g, "recommended_action": action})
    out.sort(key=lambda x: -x["volume"])
    return out[:top_n]


def conditional_opportunities(all_candidates: list[dict], in_scope_set: set[str],
                              top_n: int = 10) -> list[dict]:
    """Keywords that were classified out-of-scope but might unlock value if the
    client adds them to catalog. Top N by competitor traffic."""
    in_scope_lower = {k.lower() for k in in_scope_set}
    out_of_scope = [c for c in all_candidates if c["keyword"].lower() not in in_scope_lower]
    out_of_scope.sort(key=lambda x: -x["competitor_traffic"])
    return out_of_scope[:top_n]


def big_wins(gap_kws: list[dict], top_n: int = 3, client_name: str = "Your site") -> list[dict]:
    scored = sorted(gap_kws, key=_score_big_win, reverse=True)
    out = []
    for g in scored[:top_n]:
        if g["client_rank"] == "NR":
            pitch = (
                f"Build a dedicated page on '{g['keyword']}' (search volume {g['volume']:,}/mo). "
                f"{g['best_competitor']} currently captures {g['competitor_traffic']:,} clicks/mo from this term — "
                f"reference page structure: {g['competitor_url']}"
            )
        else:
            pitch = (
                f"Push '{g['keyword']}' from rank {g['client_rank']} into the top 5. "
                f"{g['best_competitor']} (rank {g['competitor_rank']}) captures {g['competitor_traffic']:,} clicks/mo on this term today."
            )
        out.append({**g, "pitch": pitch, "score": round(_score_big_win(g), 1)})
    return out


# ── Verdict-prose sanitiser (strips internal ROI math from client-facing text)

_ROI_LINE_RE = re.compile(
    r"(?:₹|Rs\.|INR\s|\bAOV\b|\bROI\b|\bconversion\b|\bconv\.?\s|\bretainer\b|"
    r"\bincremental revenue\b|\btimes the retainer\b|\bx\s+(?:the\s+)?retainer\b|"
    r"\bclicks/?(?:mo|month)\b.*(?:×|x)\b)",
    re.IGNORECASE,
)


def _sanitise_verdict_for_report(text: str) -> str:
    """Drop any sentence/line that contains internal ROI math from the verdict
    prose before it gets injected into the client-facing DOCX. Belt-and-braces
    safeguard — even if Claude accidentally writes AOV/retainer math in the
    'report' verdict file, the script strips it before save."""
    if not text:
        return ""
    # Split into sentences (rough — splits on period/question/exclamation +
    # whitespace, keeps clause-internal periods like "Rs.10,000" together)
    parts = re.split(r"(?<=[\.\?\!])\s+", text.strip())
    clean = [p for p in parts if not _ROI_LINE_RE.search(p)]
    out = " ".join(clean).strip()
    # If we stripped everything, fall back to a neutral placeholder
    if not out:
        return "Detailed SEO opportunity analysis follows below."
    return out


# ── Word report (compact 5-section format) ───────────────────────────────────

def write_report(out_path: Path,
                 client_name: str, client_url: str,
                 client_traffic: int,
                 comps: list[tuple[str, int]],
                 verdict_text: str,
                 gap_kws: list[dict], wins: list[dict],
                 achievable_top10: int,
                 page_opps: list[dict] | None = None,
                 quick_wins: list[dict] | None = None,
                 conditional: list[dict] | None = None,
                 misalignment: str = "",
                 niche: str = "General") -> None:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    BRAND = RGBColor(0x36, 0x75, 0x88)
    DARK  = RGBColor(0x1A, 0x1A, 0x2E)
    GREY  = RGBColor(0x6B, 0x72, 0x80)
    RED   = RGBColor(0xDC, 0x26, 0x26)
    GREEN = RGBColor(0x05, 0x96, 0x69)

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.5)
        section.left_margin = section.right_margin = Cm(1.8)

    # Compact cover — one title line + one metadata line
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(f"SEO Opportunity Audit — {client_name}")
    r.font.size = Pt(16); r.bold = True; r.font.color.rgb = BRAND

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(f"{niche}  ·  Prepared by Growisto  ·  {datetime.now().strftime('%d %b %Y')}")
    rs.font.size = Pt(10); rs.font.color.rgb = GREY

    # Verdict — large, prominent, sanitised
    doc.add_paragraph()
    vh = doc.add_paragraph()
    vr = vh.add_run("Verdict")
    vr.bold = True; vr.font.size = Pt(14); vr.font.color.rgb = BRAND
    safe_verdict = _sanitise_verdict_for_report(verdict_text)
    vp = doc.add_paragraph(safe_verdict)
    if vp.runs:
        vp.runs[0].font.size = Pt(11)

    # Misalignment / competitor-sanity warning (only if triggered)
    if misalignment:
        wp = doc.add_paragraph()
        wr = wp.add_run("⚠ Competitor selection note")
        wr.bold = True; wr.font.color.rgb = RED; wr.font.size = Pt(11)
        mp = doc.add_paragraph(misalignment)
        if mp.runs:
            mp.runs[0].font.size = Pt(10)

    # 1. Non-Branded Organic Traffic
    th = doc.add_paragraph()
    tr_run = th.add_run("1. Non-Branded Organic Traffic")
    tr_run.bold = True; tr_run.font.size = Pt(14); tr_run.font.color.rgb = BRAND
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Site"
    hdr[1].text = "Non-branded organic traffic (monthly)"
    hdr[2].text = "Gap vs " + client_name
    # Client row (bolded via run, named explicitly — no "CLIENT" label)
    row = table.add_row().cells
    for cell in row:
        for p in cell.paragraphs:
            p.clear()
    row[0].paragraphs[0].add_run(client_name).bold = True
    row[1].paragraphs[0].add_run(f"{client_traffic:,}").bold = True
    row[2].paragraphs[0].add_run("—").bold = True
    for cname, ctraf in sorted(comps, key=lambda x: -x[1]):
        row = table.add_row().cells
        row[0].text = cname
        row[1].text = f"{ctraf:,}"
        row[2].text = f"{ctraf / client_traffic:.1f}x" if client_traffic else "—"

    # (Projection callout removed — report shows facts, not estimates)

    # 2. Top Page Opportunities — single compact table, no prose
    doc.add_paragraph()
    wh = doc.add_paragraph()
    wr2 = wh.add_run("2. Top Page Opportunities")
    wr2.bold = True; wr2.font.size = Pt(14); wr2.font.color.rgb = BRAND

    if not page_opps:
        doc.add_paragraph("No page-level opportunities computed.").runs[0].font.size = Pt(10)
    else:
        pt = doc.add_table(rows=1, cols=5)
        pt.style = "Light Grid Accent 1"
        hdr = pt.rows[0].cells
        hdr[0].text = "Page topic"
        hdr[1].text = f"{client_name} clicks/mo"
        hdr[2].text = "Best competitor"
        hdr[3].text = "Competitor clicks/mo"
        hdr[4].text = "Action"
        for pg in page_opps[:5]:
            row = pt.add_row().cells
            row[0].text = pg["page_label"].title()[:40]
            row[1].text = f"{pg.get('client_traffic_on_cluster', 0):,}"
            row[2].text = pg["competitor"]
            row[3].text = f"{pg['competitor_traffic_capture']:,}"
            row[4].text = "Optimise" if pg["action_type"] == "OPTIMISE EXISTING PAGE" else "Build"

    # 3. Top Keyword Opportunities — top 10, with client rank
    doc.add_paragraph()
    kh = doc.add_paragraph()
    kr = kh.add_run("3. Top Keyword Opportunities")
    kr.bold = True; kr.font.size = Pt(14); kr.font.color.rgb = BRAND
    kt = doc.add_table(rows=1, cols=5)
    kt.style = "Light Grid Accent 1"
    hdr = kt.rows[0].cells
    hdr[0].text = "Keyword"
    hdr[1].text = "Search vol/mo"
    hdr[2].text = f"{client_name} rank"
    hdr[3].text = "Best competitor"
    hdr[4].text = "Comp clicks/mo"
    for k in sorted(gap_kws, key=lambda x: -x["competitor_traffic"])[:10]:
        row = kt.add_row().cells
        row[0].text = k["keyword"][:50]
        row[1].text = f"{k['volume']:,}"
        row[2].text = k.get("client_rank", "NR")
        row[3].text = k["best_competitor"]
        row[4].text = f"{k['competitor_traffic']:,}"

    # 4. Full Gap Keyword List — all in-scope kws (up to 30), sorted by volume
    doc.add_paragraph()
    fh = doc.add_paragraph()
    fr = fh.add_run("4. Full Gap Keyword List")
    fr.bold = True; fr.font.size = Pt(14); fr.font.color.rgb = BRAND
    ft = doc.add_table(rows=1, cols=6)
    ft.style = "Light Grid Accent 1"
    hdr = ft.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Keyword"
    hdr[2].text = "Vol/mo"
    hdr[3].text = f"{client_name} rank"
    hdr[4].text = "Best competitor"
    hdr[5].text = "Comp clicks/mo"
    for i, k in enumerate(sorted(gap_kws, key=lambda x: -x["volume"]), 1):
        row = ft.add_row().cells
        row[0].text = str(i)
        row[1].text = k["keyword"][:60]
        row[2].text = f"{k['volume']:,}"
        row[3].text = k.get("client_rank", "NR")
        row[4].text = k["best_competitor"]
        row[5].text = f"{k['competitor_traffic']:,}"

    # 5. Quick Wins — Existing Weak Rankings (pos 11-30)
    if quick_wins:
        doc.add_paragraph()
        qh = doc.add_paragraph()
        qr = qh.add_run("5. Quick Wins — Existing Weak Rankings (pos 11–30)")
        qr.bold = True; qr.font.size = Pt(14); qr.font.color.rgb = BRAND
        qt = doc.add_table(rows=1, cols=5)
        qt.style = "Light Grid Accent 1"
        hdr = qt.rows[0].cells
        hdr[0].text = "Keyword"
        hdr[1].text = "Vol/mo"
        hdr[2].text = f"{client_name} pos"
        hdr[3].text = "Comp clicks/mo"
        hdr[4].text = "Recommended action"
        for w in quick_wins:
            row = qt.add_row().cells
            row[0].text = w["keyword"][:40]
            row[1].text = f"{w['volume']:,}"
            row[2].text = w["client_rank"]
            row[3].text = f"{w['competitor_traffic']:,}"
            row[4].text = w["recommended_action"]

    # 6. Conditional Opportunities — kws rejected by in-scope filter, verify catalog
    if conditional:
        doc.add_paragraph()
        ch = doc.add_paragraph()
        cr = ch.add_run("6. Conditional Opportunities — Verify Catalog First")
        cr.bold = True; cr.font.size = Pt(14); cr.font.color.rgb = BRAND
        intro = doc.add_paragraph()
        irun = intro.add_run(
            "These keywords were classified out-of-scope (client may not currently sell these "
            "products). If catalog is expanded to include them, they unlock additional gap opportunity:"
        )
        irun.font.size = Pt(10); irun.italic = True; irun.font.color.rgb = GREY
        ct = doc.add_table(rows=1, cols=4)
        ct.style = "Light Grid Accent 1"
        hdr = ct.rows[0].cells
        hdr[0].text = "Keyword"
        hdr[1].text = "Vol/mo"
        hdr[2].text = "Best competitor"
        hdr[3].text = "Comp clicks/mo"
        for c in conditional:
            row = ct.add_row().cells
            row[0].text = c["keyword"][:50]
            row[1].text = f"{c['volume']:,}"
            row[2].text = c["best_competitor"]
            row[3].text = f"{c['competitor_traffic']:,}"
        total_add = sum(c["competitor_traffic"] for c in conditional)
        note = doc.add_paragraph()
        nrun = note.add_run(
            f"If all conditional clusters are confirmed, +{total_add:,} clicks/mo additional opportunity."
        )
        nrun.font.size = Pt(10); nrun.bold = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ── Misalignment detector ───────────────────────────────────────────────────

def detect_misalignment(client_traffic: int, comps_filtered_dfs: list[pd.DataFrame],
                        comps_total_traffic: int, in_scope_kws: set[str] | None,
                        gap_kw_count: int) -> str:
    if not in_scope_kws or not comps_total_traffic:
        return ""
    in_scope_lower = {k.lower() for k in in_scope_kws}
    relevant_traffic = 0
    for cnb in comps_filtered_dfs:
        if "Current organic traffic" in cnb.columns:
            mask = cnb["Keyword"].astype(str).str.lower().isin(in_scope_lower)
            relevant_traffic += int(cnb[mask]["Current organic traffic"].sum())
    coverage_pct = (relevant_traffic / comps_total_traffic) * 100
    # Only fire when there's almost no overlap — analyst already picked the
    # competitors, so the threshold should flag genuinely wrong picks, not
    # cases where competitors happen to be broader (gold + silver) than the
    # client (diamond-only). 19 in-scope kws is a normal narrow B2C catalog.
    if gap_kw_count < 3:
        return (f"Only {gap_kw_count} keywords passed the catalog-scope filter. "
                f"This usually means the competitor pick is wrong OR the client's "
                f"product range is too narrow to compete on these competitors' terms. "
                f"Verify before sharing.")
    if coverage_pct < 2:
        return (f"In-scope keywords represent only {coverage_pct:.1f}% of competitor traffic. "
                f"The competitor catalogs are much broader than the client's — the bulk of "
                f"their traffic is on products the client doesn't sell. The in-scope opportunity "
                f"below is real, but it's a small slice of the competitor pie.")
    return ""


# ── Main ─────────────────────────────────────────────────────────────────────

def main() -> int:
    p = argparse.ArgumentParser()
    p.add_argument("--client-name",   required=True)
    p.add_argument("--client-url",    required=True)
    p.add_argument("--client-csv",    required=True)
    p.add_argument("--competitor",    action="append", default=[],
                   help="NAME:CSV_PATH (repeatable, 1-4 times)")
    p.add_argument("--business-model", default="b2c_ecommerce")
    p.add_argument("--niche",         default="General")
    p.add_argument("--output-dir",    default="/tmp")
    p.add_argument("--in-scope-keywords", default=None,
                   help="Optional JSON list of keywords Claude marked in-scope (final pass)")
    p.add_argument("--verdict-text", default=None,
                   help="Verdict prose to inject into the DOCX (max ~600 chars)")
    p.add_argument("--verdict-file", default=None,
                   help="Read verdict prose from file (alternative to --verdict-text)")
    p.add_argument("--finalize",      action="store_true",
                   help="With --in-scope-keywords, generate the final DOCX report")
    args = p.parse_args()

    # Parse competitors
    competitors = []
    for spec in args.competitor:
        parts = spec.split(":", 1)
        if len(parts) != 2:
            print(f"Bad --competitor spec: {spec}", file=sys.stderr); return 1
        competitors.append((parts[0].strip(), parts[1].strip()))

    def _log_brand_split(label: str, df: pd.DataFrame, nb_df: pd.DataFrame) -> dict:
        total = len(df)
        brand_rows = total - len(nb_df)
        brand_pct = (brand_rows / total * 100) if total else 0
        total_traffic = int(df.get("Current organic traffic", pd.Series([0])).sum() or 0)
        nb_traffic = int(nb_df.get("Current organic traffic", pd.Series([0])).sum() or 0)
        print(
            f"[gap]   {label}: {total} total rows, {brand_rows} branded ({brand_pct:.0f}%)  "
            f"|  total traffic {total_traffic:,}  |  NON-BRAND traffic {nb_traffic:,}",
            file=sys.stderr,
        )
        # Warn if branded share is suspiciously low — CSV likely missing the
        # Branded flag → "non-brand traffic" will actually be total traffic
        if total >= 100 and brand_rows == 0:
            print(
                f"[gap]   ⚠ {label}: ZERO branded keywords detected. CSV may be "
                f"missing the Branded flag column — the 'non-brand' figures below "
                f"will actually be TOTAL traffic. Re-export from Ahrefs with brand "
                f"filters included.",
                file=sys.stderr,
            )
        return {"total_rows": total, "branded_rows": brand_rows,
                "total_traffic": total_traffic, "nonbrand_traffic": nb_traffic}

    print(f"[gap] Loading {args.client_name}...", file=sys.stderr)
    client_df = load_keywords(args.client_csv)
    client_format = client_df.attrs.get("source_format", "organic_keywords")
    print(f"[gap]   -> detected format: {client_format} ({len(client_df)} rows)", file=sys.stderr)
    client_nb = client_df[client_df["Branded"] == False].copy()
    client_traffic = int(client_nb["Current organic traffic"].sum())
    client_brand_audit = _log_brand_split(args.client_name, client_df, client_nb)

    print(f"[gap] Loading {len(competitors)} competitor(s)...", file=sys.stderr)
    comp_data = []
    comp_formats = []
    comp_brand_audits = []
    for name, csv_path in competitors:
        df = load_keywords(csv_path)
        fmt = df.attrs.get("source_format", "organic_keywords")
        comp_formats.append((name, fmt, len(df)))
        print(f"[gap]   {name}: {fmt} ({len(df)} rows)", file=sys.stderr)
        nb = df[df["Branded"] == False].copy()
        comp_data.append((name, nb, int(nb["Current organic traffic"].sum())))
        comp_brand_audits.append({"name": name, **_log_brand_split(name, df, nb)})

    # Apply intent + URL filter to all
    client_filt = _filter_intent_and_urls(client_nb, args.business_model)
    comps_filt = [(name, _filter_intent_and_urls(nb, args.business_model)) for name, nb, _ in comp_data]

    # Apply Claude's in-scope keyword filter if provided
    in_scope_set: set[str] = set()
    rejected_by_serp_guard: list[str] = []
    if args.in_scope_keywords:
        in_scope_data = json.loads(Path(args.in_scope_keywords).read_text())
        if isinstance(in_scope_data, list):
            in_scope_set = {(kw if isinstance(kw, str) else kw.get("keyword", "")).lower()
                            for kw in in_scope_data}
        elif isinstance(in_scope_data, dict) and "in_scope" in in_scope_data:
            in_scope_set = {k.lower() for k in in_scope_data["in_scope"]}

        # SERP-intent guard: when keyword starts with "<metal> <product>"
        # (e.g. "gold ring", "silver bracelet"), the SERP serves plain-metal
        # pages, not diamond-set pieces. Claude shouldn't classify these as
        # in-scope just because the client has yellow-gold-with-diamond SKUs.
        # Reject any in-scope kw matching this pattern unless the niche
        # explicitly mentions plain-metal (e.g. "22kt gold jewellery").
        niche_lower = (args.niche or "").lower()
        client_sells_plain_gold   = any(s in niche_lower for s in ["22kt gold", "22k gold", "plain gold", "gold jewellery"])
        client_sells_plain_silver = "silver" in niche_lower and "diamond" not in niche_lower.split("silver")[0]
        plain_metal_re = re.compile(r"^\s*(gold|silver|22kt|22k)\b", re.I)

        guarded: set[str] = set()
        for kw in in_scope_set:
            if plain_metal_re.search(kw):
                if "gold" in kw and not client_sells_plain_gold:
                    rejected_by_serp_guard.append(kw)
                    continue
                if "silver" in kw and not client_sells_plain_silver:
                    rejected_by_serp_guard.append(kw)
                    continue
            guarded.add(kw)
        in_scope_set = guarded

    # Keep a reference to the pre-scope competitor data — used to compute
    # conditional opportunities (candidates Claude rejected as out-of-scope)
    comps_filt_pre_scope = comps_filt

    if in_scope_set:
        comps_filt = [(name, nb[nb["Keyword"].astype(str).str.lower().isin(in_scope_set)])
                      for name, nb in comps_filt]

    print(f"[gap] Computing gap keywords...", file=sys.stderr)
    gap_kws = compute_gap_keywords(client_filt, comps_filt)
    wins = big_wins(gap_kws, client_name=args.client_name)
    page_opps = page_opportunities(gap_kws, top_n=5)
    print(f"[gap] {len(page_opps)} page-level opportunities aggregated", file=sys.stderr)

    # Quick wins (existing weak rankings, pos 11-30) — fastest action items
    quick_wins = quick_wins_existing_rankings(gap_kws, top_n=10)
    print(f"[gap] {len(quick_wins)} existing weak-ranking quick wins", file=sys.stderr)

    # Conditional opportunities (candidates rejected by in-scope filter)
    # Only computed when in_scope_set is provided (i.e. Pass 2 / --finalize)
    conditional = []
    if in_scope_set:
        all_candidates = compute_gap_keywords(client_filt, comps_filt_pre_scope,
                                              max_gap_keywords=100)
        conditional = conditional_opportunities(all_candidates, in_scope_set, top_n=10)
        print(f"[gap] {len(conditional)} conditional opportunities (rejected by scope filter)",
              file=sys.stderr)

    # Achievable target — top 10 by score
    top10 = sorted(gap_kws, key=_score_big_win, reverse=True)[:10]
    achievable = sum(g["competitor_traffic"] for g in top10)

    # Misalignment check
    comps_total = sum(t for _, _, t in comp_data)
    misalign = detect_misalignment(
        client_traffic, [df for _, df in comps_filt],
        comps_total, in_scope_set if in_scope_set else None,
        len(gap_kws),
    )

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    slug = re.sub(r"[^a-z0-9]+", "-", args.client_name.lower()).strip("-")

    # Always write the JSON dump for Claude to read
    payload = {
        "client":      {"name": args.client_name, "url": args.client_url, "traffic": client_traffic},
        "competitors": [{"name": n, "traffic": t} for n, _, t in comp_data],
        "business_model":         args.business_model,
        "niche":                  args.niche,
        "candidate_keywords":     [g["keyword"] for g in gap_kws],
        "gap_keywords":           gap_kws,
        "big_wins":               wins,
        "page_opportunities":     page_opps,
        "quick_wins":             quick_wins,
        "conditional_opportunities": conditional,
        "achievable_top10_traffic": achievable,
        "competitor_misalignment":  misalign,
        "in_scope_keywords_applied": sorted(in_scope_set) if in_scope_set else None,
        "rejected_by_serp_guard":   sorted(rejected_by_serp_guard) if rejected_by_serp_guard else None,
        "csv_formats": {
            "client": client_format,
            "competitors": [{"name": n, "format": f, "rows": r} for n, f, r in comp_formats],
        },
        "brand_audit": {
            "client": client_brand_audit,
            "competitors": comp_brand_audits,
        },
    }
    json_path = out_dir / f"gap-{slug}.json"
    json_path.write_text(json.dumps(payload, indent=2))
    print(f"\n[OK] Wrote {json_path}")

    # If finalising, also write the DOCX report
    if args.finalize:
        # Hard guard: refuse to finalise without a catalog-scope filter applied.
        # This was the root cause of "irrelevant keywords leak into the DOCX"
        # — the colleague's Claude ran --finalize without --in-scope-keywords.
        if not in_scope_set:
            print(
                "\n[ERROR] --finalize requires --in-scope-keywords. The catalog-scope\n"
                "        filter is mandatory in the final report — without it, the\n"
                "        DOCX will include out-of-scope keywords (gold/silver for a\n"
                "        diamond-only brand, etc.).\n\n"
                "        Workflow:\n"
                "          1. Run once WITHOUT --finalize to get candidate_keywords\n"
                "          2. Classify in-scope kws to /tmp/in-scope.json\n"
                "          3. Re-run with --in-scope-keywords /tmp/in-scope.json --finalize\n",
                file=sys.stderr
            )
            return 2

        # Sanity check: warn if client traffic is larger than the largest
        # competitor's non-brand traffic. That usually means the competitor
        # picks are wrong — they're smaller players than the client, so there
        # is no meaningful gap to capture from them.
        max_comp_traffic = max((t for _, _, t in comp_data), default=0)
        if client_traffic > 0 and max_comp_traffic > 0 and client_traffic > max_comp_traffic:
            warning = (
                f"⚠ COMPETITOR SELECTION WARNING: "
                f"{args.client_name} non-brand traffic ({client_traffic:,}) is LARGER "
                f"than the biggest competitor ({max_comp_traffic:,}). The chosen "
                f"competitors are smaller than the client — there is no meaningful "
                f"traffic gap to capture from them. Re-run /seo-find-competitors and "
                f"pick peers that are similar-size or larger before sharing this report."
            )
            print(f"\n{warning}\n", file=sys.stderr)
            # Surface in misalign so it appears in the DOCX
            misalign = (misalign + "\n\n" + warning).strip() if misalign else warning
            payload["competitor_misalignment"] = misalign

        verdict_text = ""
        if args.verdict_text:
            verdict_text = args.verdict_text
        elif args.verdict_file:
            verdict_text = Path(args.verdict_file).read_text(encoding="utf-8").strip()
        if not verdict_text:
            verdict_text = "(No verdict provided — pass --verdict-text or --verdict-file when finalising.)"
        report_path = out_dir / f"report-{slug}.docx"
        write_report(
            report_path, args.client_name, args.client_url, client_traffic,
            [(n, t) for n, _, t in comp_data],
            verdict_text, gap_kws, wins, achievable,
            page_opps=page_opps,
            quick_wins=quick_wins,
            conditional=conditional,
            misalignment=misalign, niche=args.niche,
        )
        print(f"[OK] Wrote {report_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
